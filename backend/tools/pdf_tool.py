"""PDF and Document generation tool with LLM-powered content."""
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional, Dict, Any
import logging
import os
from datetime import datetime

logger = logging.getLogger(__name__)


class PDFDocTool:
    """Tool for generating professional PDF and Word documents."""
    
    def __init__(self, output_dir: str = "generated_docs"):
        """
        Initialize PDF/Doc tool.
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_pdf(
        self,
        title: str,
        content: str,
        author: Optional[str] = "AI Assistant",
        subject: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a professional PDF document.
        
        Args:
            title: Document title
            content: Main content (can include sections separated by '##')
            author: Document author
            subject: Document subject
            
        Returns:
            Dictionary with file path and metadata
        """
        try:
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create PDF
            doc = SimpleDocTemplate(
                filepath,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12,
                alignment=TA_JUSTIFY,
                leading=14
            )
            
            # Build document
            story = []
            
            # Title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # Metadata
            if author or subject:
                meta_text = f"<i>Author: {author}</i>" if author else ""
                if subject:
                    meta_text += f"<br/><i>Subject: {subject}</i>"
                meta_text += f"<br/><i>Generated: {datetime.now().strftime('%B %d, %Y')}</i>"
                story.append(Paragraph(meta_text, styles['Normal']))
                story.append(Spacer(1, 0.3 * inch))
            
            # Content - parse sections
            sections = content.split('##')
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Check if section has a heading
                lines = section.split('\n', 1)
                if len(lines) > 1 and lines[0].strip():
                    # First line is heading
                    story.append(Paragraph(lines[0].strip(), heading_style))
                    section_content = lines[1].strip()
                else:
                    section_content = section
                
                # Add paragraphs
                paragraphs = section_content.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        # Handle bullet points
                        if para.startswith('•') or para.startswith('-'):
                            para = para.replace('•', '&bull;').replace('-', '&bull;')
                        story.append(Paragraph(para, body_style))
                        story.append(Spacer(1, 0.1 * inch))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF generated successfully: {filepath}")
            
            return {
                "filepath": filepath,
                "filename": filename,
                "title": title,
                "size_bytes": os.path.getsize(filepath),
                "created_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Failed to generate PDF: {e}")
            raise
    
    def generate_docx(
        self,
        title: str,
        content: str,
        author: Optional[str] = "AI Assistant",
        subject: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a professional Word document.
        
        Args:
            title: Document title
            content: Main content (can include sections separated by '##')
            author: Document author
            subject: Document subject
            
        Returns:
            Dictionary with file path and metadata
        """
        try:
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Create document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = title
            doc.core_properties.author = author or "AI Assistant"
            if subject:
                doc.core_properties.subject = subject
            
            # Title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(24)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(26, 26, 26)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            if author or subject:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if author:
                    meta_run = meta_para.add_run(f"Author: {author}\n")
                    meta_run.font.italic = True
                    meta_run.font.size = Pt(10)
                
                if subject:
                    meta_run = meta_para.add_run(f"Subject: {subject}\n")
                    meta_run.font.italic = True
                    meta_run.font.size = Pt(10)
                
                date_run = meta_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
                date_run.font.italic = True
                date_run.font.size = Pt(10)
                
                doc.add_paragraph()  # Spacer
            
            # Content - parse sections
            sections = content.split('##')
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Check if section has a heading
                lines = section.split('\n', 1)
                if len(lines) > 1 and lines[0].strip():
                    # Add heading
                    heading = doc.add_heading(lines[0].strip(), level=2)
                    heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                    section_content = lines[1].strip()
                else:
                    section_content = section
                
                # Add paragraphs
                paragraphs = section_content.split('\n\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        para = doc.add_paragraph()
                        
                        # Handle bullet points
                        if para_text.startswith('•') or para_text.startswith('-'):
                            para.style = 'List Bullet'
                            para_text = para_text.lstrip('•-').strip()
                        
                        run = para.add_run(para_text)
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"DOCX generated successfully: {filepath}")
            
            return {
                "filepath": filepath,
                "filename": filename,
                "title": title,
                "size_bytes": os.path.getsize(filepath),
                "created_at": datetime.now().isoformat()
            }
        
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            raise
    
    def get_file(self, filename: str) -> str:
        """
        Get full path to a generated file.
        
        Args:
            filename: Name of the file
            
        Returns:
            Full file path
        """
        filepath = os.path.join(self.output_dir, filename)
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filename}")
        return filepath
    
    def list_files(self) -> list:
        """
        List all generated files.
        
        Returns:
            List of file information dictionaries
        """
        files = []
        for filename in os.listdir(self.output_dir):
            filepath = os.path.join(self.output_dir, filename)
            if os.path.isfile(filepath):
                files.append({
                    "filename": filename,
                    "size_bytes": os.path.getsize(filepath),
                    "created_at": datetime.fromtimestamp(os.path.getctime(filepath)).isoformat()
                })
        return files
