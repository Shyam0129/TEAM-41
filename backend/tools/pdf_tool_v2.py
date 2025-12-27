"""
PDF Tool V2 - Professional document generation with user isolation.
"""
import os
import logging
from datetime import datetime
from typing import Optional, Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class PDFToolV2:
    """Tool for generating professional PDF and Word documents with user isolation."""
    
    def __init__(self, user_id: str, base_dir: str = "generated_docs"):
        """
        Initialize PDF tool with user-specific directory.
        
        Args:
            user_id: The unique ID of the user to isolate files
            base_dir: Base directory for all documents
        """
        self.user_id = user_id
        # Create a user-specific folder for privacy and organization
        self.output_dir = os.path.join(base_dir, user_id)
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_pdf(
        self,
        title: str,
        content: str,
        author: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate a professional PDF document."""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'Title', parent=styles['Heading1'], fontSize=24, alignment=TA_CENTER, spaceAfter=20
            )
            body_style = ParagraphStyle(
                'Body', parent=styles['Normal'], fontSize=11, alignment=TA_JUSTIFY, leading=14
            )
            
            story = [Paragraph(title, title_style)]
            
            if author:
                story.append(Paragraph(f"Author: {author}", styles['Italic']))
                story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Italic']))
                story.append(Spacer(1, 0.3 * inch))
            
            # Content parsing (simplified for v2)
            for part in content.split('\n\n'):
                if part.strip():
                    story.append(Paragraph(part.strip(), body_style))
                    story.append(Spacer(1, 0.1 * inch))
            
            doc.build(story)
            logger.info(f"PDF generated for user {self.user_id}: {filepath}")
            
            return {
                "filepath": filepath,
                "filename": filename,
                "title": title,
                "created_at": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"PDF Gen Error: {e}")
            raise

    def generate_docx(self, title: str, content: str, author: Optional[str] = None) -> Dict[str, Any]:
        """Generate a professional DOCX document."""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')[:50]
            filename = f"{safe_title}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.font.size = Pt(24)
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if author:
                auth_para = doc.add_paragraph(f"Author: {author}")
                auth_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            for part in content.split('\n\n'):
                if part.strip():
                    doc.add_paragraph(part.strip())
            
            doc.save(filepath)
            return {"filepath": filepath, "filename": filename, "title": title}
        except Exception as e:
            logger.error(f"DOCX Gen Error: {e}")
            raise
